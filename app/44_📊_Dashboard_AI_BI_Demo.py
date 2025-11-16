"""
44_Dashboard_AI_BI_Demo.py
Beschrijving:
    Deze module implementeert een interactieve Streamlit-dashboard applicatie voor onderwijsinstellingen.
    Het dashboard combineert Business Intelligence (BI), Generatieve AI en Predictive Analytics om inzichten
    te bieden in studentprestaties en risico-inschattingen.
Functionaliteiten:
    - Configuratie van de Streamlit pagina met een aangepaste layout en styling.
    - Data Upload en weergave van onderwijsdata in CSV-formaat.
    - BI-analyse per opleiding waarbij gemiddelde cijfers per vak worden berekend en gevisualiseerd met behulp
      van matplotlib en plotly.
    - Generatieve AI toepassingen:
        • Genereren van een samenvatting van de onderwijsdata via generate_ai_summary.
        • Genereren van gepersonaliseerd studieadvies via generate_personalized_feedback.
    - Voorbeeld van roosteroptimalisatie met een AI-geoptimaliseerd schema.
    - Predictive Analytics:
        • Training van een voorspellend model (RandomForestClassifier) voor het inschatten van studentrisico's.
        • Risico-inschatting voor individuele studenten op basis van hun data.
        • Uitleg genereren van het risiconiveau met behulp van AI.
Belangrijke componenten:
    - Omgeving variabelen laden met dotenv voor authenticatie.
    - Integratie met OpenAI API voor AI-gebaseerde functionaliteiten.
    - Visualisaties en analyses met behulp van matplotlib, plotly en pandas.
    - Interactieve elementen en lay-out beheer via Streamlit.
Let op:
    - Er wordt controle uitgevoerd op de aanwezigheid van benodigde kolommen in de dataset.
    - Functionaliteiten voor AI-samenvattingen en gepersonaliseerd advies zijn afhankelijk van de juist geladen en
      verwerkte data.
    - Het predictieve model wordt getraind en toegepast binnen de gebruikersinterface van Streamlit, waarbij resultaten
      direct worden weergegeven.

"""

# -------------------------------------------------------------------
# Imports
# -------------------------------------------------------------------
# Importeren van benodigde bibliotheken
import os
import streamlit as st
import pandas as pd
import plotly.express as px
from matplotlib import pyplot as plt
from openai import OpenAI

# from sklearn.model_selection import train_test_split
# from sklearn.ensemble import RandomForestClassifier
# from sklearn.metrics import accuracy_score, confusion_matrix
from dotenv import load_dotenv
from PIL import Image
import requests
import time
from dashboard_utils import *
from streamlit_extras.bottom_container import bottom
from edfutils import vandaag, tijd_verschil_microsec, tijd_nu_min, set_starttijd
from logger_decorator import logger

# from typing import Any, cast
from docx import Document


# -------------------------------------------------------------------
# Constanten en configuratie
# -------------------------------------------------------------------
# Constanten
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
image = Image.open("images/sync.jpeg")
MODEL = "gpt-4o-mini"
DATA = {}
CIJFERS = []
CIJFER = 0
risicostudenten = []
mmsamenvatting = ""
antwoord = ""
student_id = 1500
data_path = "app/static/data/2025_data.csv"
df = get_df()
vandaag_str = vandaag()
start_tijd = set_starttijd()


# -------------------------------------------------------------------
# Initialiseer de sessie status
if "antwoord" not in st.session_state:
    st.session_state.antwoord = ""
if "mmsamenvatting" not in st.session_state:
    st.session_state.mmsamenvatting = ""
if "risicostudenten" not in st.session_state:
    st.session_state.risicostudenten = []
if "df" not in st.session_state:
    st.session_state.df = df
if "dff" not in st.session_state:
    st.session_state.dff = pd.DataFrame()
if "opleiding" not in st.session_state:
    st.session_state.opleiding = "Alle"
if "klas" not in st.session_state:
    st.session_state.klas = "Alle"
if "mentor" not in st.session_state:
    st.session_state.mentor = "Alle"
if "student_id" not in st.session_state:
    st.session_state.student_id = 1500
if "student_id_predict" not in st.session_state:
    st.session_state.student_id_predict = st.session_state.student_id
if "selected_opleiding" not in st.session_state:
    st.session_state.selected_opleiding = "Alle"
if "selected_klas" not in st.session_state:
    st.session_state.selected_klas = "Alle"
if "selected_mentor" not in st.session_state:
    st.session_state.selected_mentor = "Alle"
if "opleiding_has_summary" not in st.session_state:
    st.session_state["opleiding_has_summary"] = ""
if "samenvatting_analyse" not in st.session_state:
    st.session_state["samenvatting_analyse"] = {}
if "managermessages" not in st.session_state:
    st.session_state["managermessages"] = []

risicostudenten = st.session_state.risicostudenten
print(f"Risicostudenten in sessie status: {risicostudenten if risicostudenten else 0}")


# Load the dataset
# Ensure the CSV file is in the correct path relative to this script
# @st.cache_data(ttl=3600, show_spinner=True)


def load_data():
    """
    Loads data and resources required for the dashboard.

    Returns:
        tuple: A tuple containing the following elements:
            - data_frame (pd.DataFrame): The data loaded from a CSV file
            located at "app/static/data/data.csv".
            - local_features (list of str): A list of feature names used in the dashboard.
            - dashboard_image (PIL.Image.Image): An image object
            loaded from "images/dashboard_achtergrond.png".
    """
    global df

    try:
        if df is not None:
            data_frame = df.copy()
        else:
            data_frame = pd.DataFrame()
        local_features = [
            "Burgerschap",
            "Nederlands",
            "Project_KT2",
            "Rekenen",
            "Communicatie",
            "Cijfer",
            "Aanwezigheid",
            "Waarschuwingen",
            "EC",
        ]

        dashboard_image = Image.open("images/sync.jpeg")
        return data_frame, local_features, dashboard_image
    except Exception as e:
        # Always return a tuple, even on error
        return pd.DataFrame(), [], None


# -------------------------------------------------------------------
# Data upload functionaliteit
# Laad de data en resources
df, features, image = load_data()  # type: ignore
st.session_state.df = df
st.session_state.features = features
st.session_state.image = image


# -------------------------------------------------------------------
# Hulp functies
# -------------------------------------------------------------------


def show_selection():
    """
    Displays a selection interface for filtering data based
    on 'Opleiding', 'Klas', and 'Mentor'.
    The function uses Streamlit to create a user interface with three dropdown menus:
    - 'Opleiding': Allows the user to select an educational program or choose 'Alle' (All).
    - 'Klas': Dynamically updates based on the selected 'Opleiding'
    and allows the user to select a class or 'Alle'.
    - 'Mentor': Dynamically updates based on the selected 'Opleiding' and
    allows the user to select a mentor or 'Alle'.
    The function filters a DataFrame (`df`) based on the selected values
    from the dropdown menus. If no DataFrame is
    available, it returns an empty DataFrame.
    Returns:
        tuple: A tuple containing:
            - filtered_data (pd.DataFrame): The filtered DataFrame based on
            the selected criteria.
            - selected_opleiding (str): The selected educational program.
            - selected_klas (str): The selected class.
            - selected_mentor (str): The selected mentor.
    """
    st.markdown("##### 🎯**Kies een opleiding, klas en/of mentor**")
    ccol1, ccol2, ccol3 = st.columns(3)
    with ccol1:
        selected_opleiding = st.selectbox("**OPLEIDING**", OPLEIDINGEN + ["Alle"])

    with ccol2:
        if selected_opleiding == "Alle":
            selected_klas = st.selectbox("**KLAS**", list(KLASSEN["Alle"]))
        else:
            selected_klas = st.selectbox("**KLAS**", list(KLASSEN[selected_opleiding]))

    with ccol3:
        if selected_opleiding == "Alle":
            selected_mentor = st.selectbox("**MENTOR**", list(MENTOREN["Alle"]))
        else:
            selected_mentor = st.selectbox(
                "**MENTOR**", list(MENTOREN[selected_opleiding])
            )

    filtered_data = df.copy() if df is not None else pd.DataFrame()
    if selected_opleiding != "Alle":
        filtered_data = filtered_data[filtered_data["Opleiding"] == selected_opleiding]
    if selected_klas != "Alle":
        filtered_data = filtered_data[filtered_data["Klas"] == selected_klas]
    if selected_mentor != "Alle":
        filtered_data = filtered_data[filtered_data["Mentor"] == selected_mentor]

    if filtered_data is not None:
        return filtered_data, selected_opleiding, selected_klas, selected_mentor
    else:
        return pd.DataFrame(), "Alle", "Alle", "Alle"


# -------------------------------------------------------------------
# Configuratie van de pagina
# -------------------------------------------------------------------
# Configuratie van de pagina
st.set_page_config(
    layout="wide",
    initial_sidebar_state="collapsed",
    menu_items=None,
    page_icon="⛳",
    page_title="Dashboard GenBI in het Onderwijs",
)

col1, col2 = st.columns([1, 1])
with col1:
    st.markdown(
        """
                #### :blue[**Synergie van AI en BI in het onderwijs**]\n
                ## 📊:blue[**AI Driven BI Dashboard**]"""
    )
    with st.expander("💡 :blue[**Hoe werkt het?**]", expanded=False):
        st.markdown(
            """
            **In dit Dashboard laten we zien hoe we Generatieve AI en Business Intelligence (BI) combineren om inzichten te verkrijgen in studentprestaties en risico-inschattingen.**
                 
1. **BI-inzichten**: Data over studentvoortgang, feedback, en KPI's.     
2. **Generatieve AI Toepassingen**: Voorbeelden zoals gepersonaliseerde 
        adviezen, concept e-mails, of samenvattingen.     
3. **Predictive Analytics**: Voorspellende modellen voor studentrisico's en
        aanbevelingen voor interventies.
4. **Roosteroptimalisatie**: AI-geoptimaliseerde roosters voor betere planning.
5. **Interactieve Visualisaties**: Grafieken en analyses voor inzicht in studentprestaties.
6. **Persoonlijke Leeradviezen**: Gepersonaliseerde feedback en advies voor studenten.
7. **Doelstellingen**: Inzicht in studentprestaties, gepersonaliseerde leerervaringen, en efficiënte rapportage en feedback.
        
:green[**Start met het uploaden van een CSV-bestand met onderwijsdata**]
             
De data moet deze kolommen bevatten: **`Student-ID`, `Opleiding`, `Klas`, `Mentor`,`Burgerschap`, `Nederlands`, `Project_KT2`, `Rekenen`, `Communicatie`, `Cijfer`, `Aanwezigheid`, `Waarschuwingen`, `EC`.**

     
        """
        )
    # -------------------------------------------------------------------
    # Selectie van opleiding, klas en mentor
    # -------------------------------------------------------------------
    # Laat de gebruiker opleiding, klas en mentor selecteren, en plaats hetgefilterde      dataframe in dff
    dff, opleiding, klas, mentor = (
        show_selection() if df is not None else (pd.DataFrame(), "Alle", "Alle", "Alle")
    )
    # -------------------------------------------------------------------
    # Sla de geselecteerde waarden op in de sessie status
    st.session_state.dff = dff
    st.session_state.opleiding = opleiding
    st.session_state.klas = klas
    st.session_state.mentor = mentor

with col2:
    if image is not None:
        st.image(
            image,
            caption=None,
            width=320,
            clamp=True,
            channels="RGB",
            output_format="auto",
        )
    else:
        st.warning("Afbeelding kon niet worden geladen.")


# -------------------------------------------------------------------
# Hulp functies
# -------------------------------------------------------------------


# Download docx bestand
@st.fragment()
def download_report_docx(docpath: str = "samenvatting_analyse_rapport.docx"):
    """Download a docx file"""
    # global button_teller
    # button_teller += 1

    with open(docpath, "rb") as file:
        bestand = file.read()
        btn = st.download_button(
            label=":red[DOWNLOAD HET RAPPORT]",
            data=bestand,
            file_name="managementsamenvatting_analyse_rapport.docx",
            mime="application/docx",
            # key=str(button_teller) + "download_button_" + str(time.asctime),
        )


# -------------------------------------------------------------------
# Header en metadata
# -------------------------------------------------------------------
st.markdown(
    f"""
    <style>
    .stApp {{
        background-image: url("http://edflaptop.dyndns-home.com:8501/app/static/images/achtergrond_lichtblauw.jpeg");
        background-size: 120%;
        transparant: 100%;
        background-position: top left;
        background-repeat: no-repeat;
        background-attachment: local;
        background-color: #f0f2f5;

    }}
    </style>
    """,
    unsafe_allow_html=True,
)


st.markdown(
    f"""###### 🔶**Opleiding:** :blue[**{opleiding.strip() if opleiding.strip() != "Alle" else 'alle opleidingen'}**],   🔶**Klas:** :blue[**{klas.strip() if klas.strip() != "Alle" else 'alle klassen'}**],   🔶**Mentor:** :blue[**{mentor if mentor != "Alle" else 'alle mentoren'}**],   🔶**Aantal studenten:** :blue[**{len(dff)}**]"""
)


# -------------------------------------------------------------------


st.dataframe(
    dff,
    key="studenten_overzicht",
    # on_change=change_data,
    column_config={
        "Student-ID": st.column_config.NumberColumn(
            label="Student-ID", format="%d", help="Unieke ID van de student"
        ),
        "Naam": st.column_config.TextColumn(label="Naam", help="Naam van de student"),
        "Opleiding": st.column_config.TextColumn(
            label="Opleiding", help="Naam van de opleiding"
        ),
        "Klas": st.column_config.TextColumn(label="Klas", help="Klas van de student"),
        "Mentor": st.column_config.TextColumn(
            label="Mentor", help="Mentor van de student"
        ),
        "Burgerschap": st.column_config.NumberColumn(
            label="Burgerschap", format="%.2f", help="Cijfer voor Burgerschap"
        ),
        "Nederlands": st.column_config.NumberColumn(
            label="Nederlands", format="%.2f", help="Cijfer voor Nederlands"
        ),
        "Project_KT2": st.column_config.NumberColumn(
            label="Project KT2", format="%.2f", help="Cijfer voor Project KT2"
        ),
        "Rekenen": st.column_config.NumberColumn(
            label="Rekenen", format="%.2f", help="Cijfer voor Rekenen"
        ),
        "Communicatie": st.column_config.NumberColumn(
            label="Communicatie", format="%.2f", help="Cijfer voor Communicatie"
        ),
        "Cijfer": st.column_config.NumberColumn(
            label="Cijfer", format="%.2f", help="Gemiddeld cijfer"
        ),
        "Aanwezigheid": st.column_config.ProgressColumn(
            label="Aanwezigheid",
            format="%.1f%%",
            min_value=0,
            max_value=100,
            help="Aanwezigheid percentage",
        ),
        "Waarschuwingen": st.column_config.NumberColumn(
            label="Waarschuwingen", format="%d", help="Aantal waarschuwingen"
        ),
        "EC": st.column_config.ProgressColumn(
            label="EC",
            format="%d",
            min_value=0,
            max_value=60,
            help="Aantal behaalde EC's",
        ),
    },
)

# -------------------------------------------------------------------
# Header met metadata
# -------------------------------------------------------------------
mentor = st.session_state.mentor
klas = st.session_state.klas
opleiding = st.session_state.opleiding


# -------------------------------------------------------------------
# Statistieken en KPI's
# -------------------------------------------------------------------
col1, col2, col3, col4 = st.columns([2, 2, 2, 2])


col1.metric(
    label="###### 🔶 **CIJFER**\n\n:green[**(norm: 6.0)**]",
    value=f"✨{dff['Cijfer'].astype(float).mean():.2f}",
    delta=f"{dff['Cijfer'].astype(float).mean() - 6.0:.2f}",
    delta_color="normal",
    border=False,
)
col2.metric(
    label="###### 🔶 **AANWEZIGHEID**\n\n:green[**(norm: 85)**] ",
    value=f"✨{dff['Aanwezigheid'].mean():.1f}%",
    delta=f"{dff['Aanwezigheid'].mean()-85:.1f}%",
    delta_color="normal",
    border=False,
)
col3.metric(
    label="###### 🔶 **WAARSCHUWING**\n\n:green[**(norm: 0.5)**] ",
    value=f"✨{dff['Waarschuwingen'].mean():.2f}",
    delta=f"{0.5 - dff['Waarschuwingen'].mean():.2f}",
    delta_color="normal",
    border=False,
)
col4.metric(
    label="###### 🔶 **EC**\n\n:green[**(norm: 45)**] ",
    value=f"✨{dff['EC'].mean():.2f}",
    delta=f"{dff['EC'].mean()-40:.2f}",
    delta_color="normal",
    border=False,
)
st.write("---")
# -------------------------------------------------------------------
# --- Functies voor GenAI ---
# -------------------------------------------------------------------


def download_analyse_docx(docpath: str = "managementsamenvatting.docx"):
    """"""
    with open(docpath, "rb") as file:
        bestand = file.read()

        st.download_button(
            label=":red[DOWNLOAD DE SAMENVATTING]",
            data=bestand,
            file_name="managementsamenvatting.docx",
        )


def save_as_docx(minutes: str, filename: str, opleiding: str = opleiding):
    """_summary_

    Args:
        minutes (str): _description_
        filename (str): _description_
        opleiding (str, optional): _description_. Defaults to opleiding.

    Returns:
        _type_: _description_
    """

    doc = Document()
    doc.add_heading(f"Managementsamenvatting {opleiding}\n" + str({vandaag()}), 0)
    doc.add_paragraph()

    # minutes = minutes.encode("utf-8")

    paragraaf = ""

    for key in minutes.split():
        key = str(key) + " "
        if key.startswith("###") or key.startswith("###") or "###" in key:
            heading_level = 3
        elif key.startswith("##") or key.startswith("##") or "##" in key:
            heading_level = 2
        elif key.startswith("#") or key.startswith("#") or "#" in key:
            heading_level = 1
        else:
            heading_level = 0

        if heading_level > 0:
            if len(paragraaf) > 0:
                # Voeg de huidige paragraaf toe aan het document
                doc.add_paragraph(
                    paragraaf.replace("\\n", " ").replace("b'", " ").replace("'", " ")
                )
                doc.add_paragraph()
                paragraaf = " "

            paragraaf += (
                key.replace("###", " ")
                .replace("##", " ")
                .replace("#", " ")
                .replace("**", " ")
                .replace("*", " ")
                .replace("b'", " ")
                .replace("'", " ")
            )

            doc.add_heading(paragraaf, level=heading_level)
            doc.add_paragraph()  # Voeg een lege regel toe na de heading

            # Reset de paragraaf voor de volgende heading
            if paragraaf.strip() == "":
                # Als de paragraaf leeg is, voeg een lege regel toe
                doc.add_paragraph()
            else:
                # Voeg de paragraaf toe aan het document
                doc.add_paragraph(
                    paragraaf.replace("\\n", " ").replace("b'", " ").replace("'", " ")
                )

            paragraaf = " "
        else:

            paragraaf += str(key) + " "

            # paragraaf += paragraaf.strip()
            if paragraaf == "\\n" or "\\n" in paragraaf:
                # Voeg de paragraaf toe aan het document
                doc.add_paragraph(
                    paragraaf.replace("\\n", " ").replace("b'", " ").replace("'", " ")
                )
                doc.add_paragraph()
                paragraaf = ""
                # Voeg een lege regel toe na de paragraaf
            else:
                paragraaf += paragraaf

    doc.add_paragraph()
    doc.save(filename)

    print(f"Document opgeslagen als {filename}")

    return "doc bewaard in " + filename


def make_download(demarkdown: str, filename: str, opleidingnaam: str = opleiding):
    """
    Maak een downloadbare markdown file van de samenvatting.
    """
    # Maak een markdown bestand van de samenvatting
    opslaan_als_word = save_as_docx(demarkdown, filename, opleidingnaam)

    if opslaan_als_word:
        download_analyse_docx(docpath=filename)


def toon_prognose():
    # Laad de data
    bestandspad = "app/static/data/prognose_instroom.csv"
    opleiding = st.session_state.opleiding.upper()
    opleiding_kolom = None
    df_prognose = pd.read_csv(bestandspad)

    jaar_kolom = "Jaar"
    jaar_kolom = df_prognose.columns[0]
    # Alle kolommen behalve de eerste (Jaar)
    opleidingen = df_prognose.columns[1:]
    for i in range(len(opleidingen)):
        if opleidingen[i].upper() == opleiding.upper():
            opleiding_kolom = opleidingen[i]
            break
        else:
            if opleidingen[i].upper() in opleiding.upper():
                opleiding_kolom = opleidingen[i]
                break
            else:
                if opleiding.upper() in opleidingen[i].upper():
                    opleiding_kolom = opleidingen[i]
                    break

    if not opleiding_kolom:
        st.error(f"De opleiding '{opleiding}' is niet gevonden in de prognose data.")
        return

    prognose_opleiding_kolom = df_prognose[opleiding_kolom]
    # Toon de ruwe data (optioneel)
    # Zorg dat de jaarkolom als string of integer wordt gelezen voor correcte x-as
    df_prognose[jaar_kolom] = df_prognose[jaar_kolom].astype(str)

    # df_prognose["Jaar"] = jaar_kolom  # Update de kolom in de DataFrame

    # Lijngrafiek weergeven
    # st.markdown(f"**Prognose instroom per opleiding {opleiding}**")

    # Keuzemenu voor kolommen
    kolom = opleiding_kolom

    # df_prognose[opleiding] = df_prognose[opleiding].astype(float)  # Zorg dat de kolom als integer wordt gelezen voor correcte y-as
    # df_prognose[kolom] = df_prognose[kolom].to_numpy(dtype=float) / len(st.session_state.df["Naam"]) * 100

    # Plotten
    fig, ax = plt.subplots(squeeze=True, figsize=(12, 2), sharex="all", sharey="all")
    ax.plot(
        df_prognose[jaar_kolom],
        df_prognose[kolom],
        "D",
        linestyle="-",
        label=kolom,
        linewidth=1,
        markersize=4,
    )

    ax.set_xlabel(
        "Cohort",
    )
    ax.set_ylabel("Aantal studenten")
    ax.set_title(f"Prognoses instroom {opleiding} - {vandaag()}")
    ax.legend()
    # ax.set_xticks(df[jaar_kolom])  # Zet de x-ticks op de jaarkolom
    ax.grid(True)

    # Toon de grafiek
    st.pyplot(fig, use_container_width=True, clear_figure=True)


# ---------------------------------------------------------
# HOOFDPROGRAMMA
# ---------------------------------------------------------
# Laad de data en controleer op vereiste kolommen
#
if df is not None:
    # Controleer of de vereiste kolommen aanwezig zijn
    required_columns = ["Student-ID", "Opleiding", "Klas"] + VAKKEN + ["Cijfer"]
    missing_columns = [col for col in required_columns if col not in df.columns]
    if missing_columns:
        st.error(
            f"De volgende kolommen ontbreken in de dataset: {', '.join(missing_columns)}"
        )
        df = None  # Reset df als er ontbrekende kolommen zijn
    else:
        print("Data succesvol geladen!")

    st.subheader("📊 :blue[**Overzicht**]")
    st.markdown(
        "**Diverse statistieken mbt de studentenpopulatie en voortgang van de studenten per opleiding, klas en mentor.**"
    )
    with st.expander(
        f" :red[**Tip: Bekijk overzicht KPI's mbt selectie -**] :blue[**{opleiding.capitalize()}, {klas if klas != 'Alle' else 'alle klassen'} - {mentor if mentor != 'Alle' else 'alle mentoren'}**]",
        expanded=False,
        icon="💡",
    ):
        st.subheader(":blue[**Strategie en Governance**]")

        # Histogram en boxplot voor cijferverdeling en aanwezigheid
        # in 2 kolommen
        toon_prognose()
        st.subheader(":blue[**Trendgrafieken & spreiding**]")
        histogram_col1, boxplot_col2 = st.columns(2)
        with histogram_col1:
            fig = px.histogram(
                dff,
                x="Cijfer",
                color="EC",
                facet_col="Opleiding",
                title="Cijfer en EC verdeling ",
                nbins=len(dff),
            )
            st.plotly_chart(fig,use_container_width=True)

        with boxplot_col2:
            fig2 = px.box(
                dff,
                x="Opleiding",
                y="Aanwezigheid",
                points="all",
                title="Aanwezigheid per Opleiding",
            )

            st.plotly_chart(fig2, use_container_width=True)

        # Trendgrafieken voor gemiddelde cijfers per vak
        ccol1, ccol2 = st.columns([1, 1])

        with ccol1:
            if df is not None:
                grouped_df = df.groupby("Opleiding")["Cijfer"].mean().reset_index()
                fig3 = px.bar(
                    grouped_df,
                    x="Cijfer",
                    y="Opleiding",
                    title="Gemiddelde Cijfers per Opleiding",
                    color="Cijfer",
                    color_continuous_scale=px.colors.sequential.Blues,
                    orientation="h",
                )
                st.plotly_chart(
                    fig3,
                    height=400,
                    width=500,
                    theme="streamlit",
                )

                # st.bar_chart(
                #     grouped_df,
                #     x="Opleiding",
                #     y="Cijfer",
                #     height=400,
                #     width=500,
                #     color="Cijfer",
                #     horizontal=True,
                #     width='stretch',
                # )

            else:
                st.error("Dataframe 'df' is not loaded")
        with ccol2:

            x1 = dff["Burgerschap"].to_list()
            x2 = dff["Communicatie"].to_list()
            x3 = dff["Nederlands"].to_list()
            x4 = dff["Project_KT2"].to_list()
            x5 = dff["Rekenen"].to_list()

            # Group data together
            hist_data = [x1, x2, x3, x4, x5]
            group_labels = [
                "Burgerschap",
                "Communicatie",
                "Nederlands",
                "Project_KT2",
                "Rekenen",
            ]

            mean_values = dff[VAKKEN].mean().tolist()
            fig1 = px.pie(
                dff[VAKKEN],
                names=group_labels,
                values=mean_values,
                template="plotly_white",
                color_discrete_sequence=[
                    "#0068c9",
                    "#83c9ff",
                    "#ff2b2b",
                    "#ffabab",
                    "#29b09d",
                ],
                title=f"EC verdeling per vak:\n\n{opleiding.capitalize()} - {klas.lower()}",
            )

            st.plotly_chart(
                fig1,
                width=250,
                height=100,
                use_container_width=True,
                theme="streamlit",
                key="pie_chart_1",
            )

        if df is not None:
            df_opleiding = (
                df[df["Opleiding"] == opleiding].copy()
                if df is not None
                else pd.DataFrame()
            )
        else:
            st.error("Geen data beschikbaar. Upload eerst een CSV-  bestand.")
            st.stop()

        # st.write(
        #     f":blue[**{opleiding.capitalize()} klassen benchmark**]"
        # )
        if df is not None:
            if df is not None:
                if df is not None:
                    if df is not None:
                        if df is not None:
                            if df is not None:
                                if df is not None:
                                    df_opleiding = df[
                                        df["Opleiding"] == opleiding
                                    ].copy()
                                else:
                                    df_opleiding = (
                                        pd.DataFrame()
                                    )  # Initialize as an empty DataFrame
                            else:
                                st.error("Dataframe 'df' is not loaded or is None.")
                                df_opleiding = (
                                    pd.DataFrame()
                                )  # Initialize as an empty DataFrame
                        else:
                            st.error("Dataframe 'df' is not loaded or is None.")
                            df_opleiding = (
                                pd.DataFrame()
                            )  # Initialize as an empty DataFrame
                    else:
                        st.error("Dataframe 'df' is not loaded or is None.")
                        df_opleiding = (
                            pd.DataFrame()
                        )  # Initialize as an empty DataFrame
                else:
                    st.error("Dataframe 'df' is not loaded or is None.")
                    df_opleiding = pd.DataFrame()  # Initialize as an empty DataFrame
            else:
                st.error("Dataframe 'df' is not loaded or is None.")
                df_opleiding = pd.DataFrame()  # Initialize as an empty DataFrame
        else:
            st.error("Dataframe 'df' is not loaded")
            df_opleiding = pd.DataFrame()  # Initialize as empty DataFrame

        grouped_df = df_opleiding.groupby("Klas")[VAKKEN].mean().reset_index()
        grouped_df["Klas"] = grouped_df["Klas"].str.upper()
        grouped_df = grouped_df.sort_values(by="Klas")
        fig_bar = px.bar(
            grouped_df,
            x="Klas",
            y=VAKKEN,
            barmode="group",
            title="Benchmark klassen - gemiddelde cijfers per vak",
        )
        st.plotly_chart(fig_bar, use_container_width=True)

        fig7 = px.line(
            grouped_df,
            x="Klas",
            y=VAKKEN,
            title=f"Gem. cijfer per vak voor {klas.lower()}",
            markers=True,
        )
        fig7.update_layout(
            title=f"Gem. cijfer per vak voor klas {klas.lower()}",
            xaxis_title="Klas",
            yaxis_title="Gemiddelde Cijfer",
            legend_title="Vakken",
        )
        st.plotly_chart(fig7, use_container_width=True, height=400, width=500)

        if df is not None:
            df_opleiding = df[df["Opleiding"] == opleiding].copy()
        else:
            df_opleiding = pd.DataFrame()
        # st.write(f":blue[**{opleiding.capitalize()} - alle mentoren - cijfer per vak**]")
        fig_mark = px.bar(
            df_opleiding,
            x="Mentor",
            y=VAKKEN,
            barmode="group",
            title="Gemiddeld cijfer per vak bij mentor",
        )
        fig_mark.update_layout(
            xaxis_title="Mentor",
            yaxis_title="Gemiddelde cijfer per vak bij mentor",
            height=450,
            width=670,
            legend_title="Vakken",
            xaxis_tickangle=-45,
        )
        st.plotly_chart(fig_mark, )

    # -------------------------------------------------------------------
    # Genereer AI-samenvatting
    @st.fragment()
    def choose(opleiding: str = opleiding):
        """
        Kies een opleiding voor samenvatting.
        """
        global df
        global antwoord
        if df is None:
            st.error("Geen data geladen. Upload eerst een CSV-bestand.")
            st.stop()

        # st.write("### Selecteer Opleiding voor AI-Samenvatting")
        if opleiding == "Alle":
            st.info("Selecteer een specifieke opleiding voor samenvatting.")
            return "Geen specifieke opleiding geselecteerd."

        df_opleiding = df[df["Opleiding"] == opleiding].copy()
        start_tijd = set_starttijd()
        st.write(tijd_verschil_microsec(start_tijd))

        with st.spinner("Bezig met het analyseren van de data...", show_time=True):

            prompt = f"""
                Jij bent een onderwijsadviseur en data-analist je maakt uitgebreide management samenvattingen op basis van data-analyse op een data-set.
                Je gaat nu een uitgebreide samenvatting maken van de bijgevoegde data van de opleiding {opleiding}.
                Je analyseert de data van de studenten, en maakt een samenvatting van de prestaties van de studenten in de opleiding {opleiding}.
             
                Hiervoor gebruik je bijgevoegde data van de opleiding: {opleiding}, de vakken {VAKKEN}, de klassen {KLASSEN[opleiding]}, en de mentoren {MENTOREN[opleiding]}. Je gebruikt uit de data de kolommen {VAKKEN}, Cijfer, Aanwezigheid, en EC. EC staat voor ECTS (European Credit Transfer and Accumulation System) en is een maat voor studiepunten die studenten behalen in het MBO-onderwijs. In het onderwijs worden EC's gebruikt om de studiebelasting van een opleiding aan te geven. Een EC staat voor 28 uur studiebelasting, inclusief contacturen, zelfstudie, en examens. Studenten moeten een bepaald aantal EC's behalen om hun opleiding succesvol af te ronden. In het MBO-onderwijs zijn er vaak specifieke eisen voor het aantal EC's dat studenten moeten behalen in verschillende vakken of modules. Om een positieve beoordeling te krijgen, moeten studenten voldoen aan de norm van 40 EC's in het eerste jaar. 
                
                Naast de studiepunten zijn ook de cijfers voor de vakken, het gemiddelde Cijfer, het aanwezigheidspercentage, moet minimaal 85% zijn, en het aantal waarschuwingen, moet kleiner dan 1 zijn, bepalend voor het Bindend Studie Advies.
                EC zijn behaalde studiepunten en bepalen mede of een student door mag naar het tweede jaar. De norm is op dit moment 40 punten. 
                  
                Werk stap voor stap. Analyseer eerst de data, en rapporteer daarna je bevindingen (wat gaat goed, wat gaat niet goed, opvallende zaken, trends, enz.) in een uitgebreide samenvatting. De title is:\n
                
                <titel>
                Managementsamenvatting voor {opleiding} - {vandaag()}
                </titel>
                
                Gebruik de volgende structuur voor de samenvatting:
                1. **Inleiding**: Geef een korte inleiding over de opleiding en de data.
                2. **Analyse**: Analyseer de data en geef een overzicht van de prestaties per klas, vak, en mentor.
                3. **Conclusie**: Geef een samenvatting van de belangrijkste bevindingen.
                4. **Aanbevelingen**: Geef aanbevelingen voor verbetering op basis van de analyse.
                5. **Risicostudenten**: Identificeer risicostudenten op basis van de data en geef aanbevelingen voor ondersteuning.
                6. **Bindend Studie Advies**: Geef een overzicht van de studenten die mogelijk een Bindend Studie Advies (BSA) nodig hebben op basis van de data.
                
                7. **Download optie**: Zorg ervoor dat de samenvatting in een docx bestand wordt geplaatst, en biedt dat als download optie aan.
                8. **Markdown formaat**: Gebruik het Markdown formaat voor de samenvatting en antwoord altijd in het Nederlands.
                
                
                
                Zorg er tenslotte voor dat je de samenvatting en al je bevindigen in een docx bestand plaatst, en biedt die als download optie aan. 
                Hieronder de data die je gaat analyseren:\n\n
                
                <data>
                
                {df_opleiding.to_string(index=False)}
                
                </data>
                
                """

            resp = requests.post(
                "http://edflaptop.dyndns-home.com:8000/analyseer",
                headers={"Content-Type": "application/json"},
                json={"data": prompt},
                timeout=220,
            )
            # antwoord = "
            antwoord = resp.json()["summary"]

        if antwoord is not None and len(antwoord) > 0:
            st.markdown(f"### AI-Samenvatting voor {opleiding}")
            st.info(antwoord)
            st.write(tijd_verschil_microsec(start_tijd))

            return antwoord

        if antwoord is None:
            return None

    # -------------------------------------------------------------------

    # st.write("----------------------------------------")
    # GenAI Rapport
    st.subheader("🎯 :blue[**AI Management rapportages**]")
    st.markdown(
        """
        **Laat de AI een analyse uitvoeren op alle studenten van de geselecteerde sector.
        De resultaten komen in een te downloaden rapport**"""
    )
    st.markdown(
        "💡**Let op: dit kan even duren (gemiddeld 1 minuut), afhankelijk van de grootte van de dataset.**"
    )

    st.fragment()

    def choose_opleiding_summary():
        """
        Kies een opleiding voor samenvatting.
        """
        global df
        global antwoord
        global opleiding
        rapportage_opleiding = opleiding
        summary_gemaakt = None

        # opleiding = st.session_state.opleiding
        if rapportage_opleiding == "Alle":
            st.info("Selecteer een specifieke opleiding voor samenvatting.")
            st.markdown(
                "##### :blue[Selecteer een Sector/Opleiding voor de Samenvatting]"
            )
            rapportage_opleiding = st.selectbox(
                ":blue[**Selecteer een sector/opleiding**]",
                options=OPLEIDINGEN,
                index=0,
            )

        if rapportage_opleiding != "Alle":
            st.session_state.opleiding = rapportage_opleiding
            if (
                "samenvatting_analyse" in st.session_state
                and st.session_state["samenvatting_analyse"] is not None
            ):
                if rapportage_opleiding in list(
                    st.session_state["samenvatting_analyse"].keys()
                ):
                    summary_gemaakt = st.session_state["samenvatting_analyse"][
                        rapportage_opleiding
                    ]
                    print(summary_gemaakt)
                    st.markdown("### AI-Samenvatting")
                    st.info(summary_gemaakt)

                    return summary_gemaakt

            if rapportage_opleiding != "Alle":
                if df is not None:
                    df_opleiding = df[df["Opleiding"] == rapportage_opleiding].copy()
                else:
                    st.error("Dataframe 'df' is not loaded or is None.")
                    df_opleiding = pd.DataFrame()  # Initialize as an empty DataFrame
                # st.session_state.opleiding = opleiding
                if df_opleiding.empty:
                    st.warning(f"Geen data gevonden voor opleiding: {opleiding}")
                    return "Geen data beschikbaar voor deze opleiding."
                    # st.subheader("AI-Samenvattingen per Opleiding")

                start_tijd = set_starttijd()
                st.write(tijd_verschil_microsec(start_tijd))

                with st.spinner(
                    f"Bezig met het genereren van samenvatting voor sector{rapportage_opleiding}... ",
                    show_time=True,
                ):
                    summary_gemaakt = choose(rapportage_opleiding)
                    if summary_gemaakt is not None:

                        # st.markdown("### AI-Samenvatting")
                        # st.info(summary_gemaakt)
                        if "samenvatting_analyse" not in st.session_state:
                            st.session_state["samenvatting_analyse"][
                                rapportage_opleiding
                            ] = summary_gemaakt
                        else:
                            st.session_state["samenvatting_analyse"][
                                rapportage_opleiding
                            ] = summary_gemaakt
                        if "has_summary" not in st.session_state:
                            st.session_state["opleiding_has_summary"] = (
                                rapportage_opleiding
                            )
                        else:
                            st.session_state["opleiding_has_summary"] = (
                                rapportage_opleiding
                            )
                        # data = summary_gemaakt.encode("utf-8")

                        st.write(tijd_verschil_microsec(start_tijd))

                        return summary_gemaakt
                    else:
                        return None

            return None

    # -------------------------------------------------------------------
    # Genereer samenvatting voor de geselecteerde opleiding

    with st.expander(
        f":red[**Genereer een AI-samenvatting voor de geselecteerde opleiding -**] :blue[**{opleiding.capitalize()}, {'alle klassen' if klas == 'Alle' else klas}, {'alle mentoren' if mentor == 'Alle' else mentor}**]",
        expanded=False,
        icon="💡",
    ):
        st.markdown(
            "### Genereer een AI-samenvatting van de onderwijsdata voor de geselecteerde opleiding"
        )
        st.markdown(
            "💡 **Let op: dit kan even duren (gemiddeld 1 minuut), afhankelijk van de grootte van de dataset.**"
        )
        if (
            "samenvatting_analyse" in st.session_state
            and st.session_state["samenvatting_analyse"] is not None
            and opleiding in st.session_state["samenvatting_analyse"].keys()
        ):
            summary_gemaakt = st.session_state["samenvatting_analyse"][opleiding]
        else:
            summary_gemaakt = choose_opleiding_summary()

    if summary_gemaakt:
        if "samenvatting_analyse" not in st.session_state:
            st.session_state["samenvatting_analyse"] = {}
        st.session_state["samenvatting_analyse"][opleiding] = summary_gemaakt
        st.info(f"Laatst gemaakte samenvatting: **{opleiding}** op  {vandaag()}")

    if (
        "samenvatting_analyse" in st.session_state
        and st.session_state["samenvatting_analyse"] is not None
    ):
        for samenv in st.session_state["samenvatting_analyse"].keys():

            with st.expander(
                f":green[**AI-Samenvatting voor {samenv}**]",
                expanded=False,
                icon="💡",
            ):

                # Display the AI-generated summary
                st.markdown(f"**Samenvatting voor {samenv} - {vandaag()}**")
                st.info(st.session_state["samenvatting_analyse"][samenv])

                # Download button for the summary

                st.download_button(
                    label="Download Samenvatting",
                    data=st.session_state["samenvatting_analyse"][samenv].encode(
                        "utf-8"
                    ),
                    file_name=f"samenvatting_{samenv}.md",
                    mime="text/markdown",
                )

    # -------------------------------------------------------------------
    # Voorspellende analyse
    # -------------------------------------------------------------------
    # Voorspellende analyse: Uitval van studenten
    @st.fragment()
    def details_uitval(Risicostudenten: list = []):
        """
        Toon details van risicostudenten en hun voorspelde uitvalrisico.
        """
        global dff
        global features
        risico_studenten_lijst = []
        st.markdown("#### :red[Details van risicostudenten]")
        for student in Risicostudenten:
            risico_studenten_lijst.append(student[0]["Naam"])

        keuze_student = st.selectbox(
            ":blue[**Selecteer een student om details te bekijken**]",
            options=["Selecteer een student"] + risico_studenten_lijst,
            index=0,
            format_func=lambda x: (
                str(x) if x is not None else "Geen student geselecteerd"
            ),
        )
        if keuze_student != "Selecteer een student":
            # Filter de DataFrame op basis van de geselecteerde student
            student_row = dff[
                dff["Naam"].str.strip() == keuze_student.split(" (")[0]
            ].iloc[0]
            # student_row = dff[dff["Naam"] == keuze_student].iloc[0]
            # Fetch prediction result for the student
            start_tijd = set_starttijd()

            pred_response = requests.post(
                "http://edflaptop.dyndns-home.com:8000/predict_dropout",
                json={"student": student_row[features].to_dict()},
                timeout=60,
            )
            result = pred_response.json()
            # st.write(tijd_verschil_microsec(start_tijd))
            # st.write(result)
            # Display the prediction result
            st.markdown(
                f"🔴 :red[**{student_row['Student-ID']},{student_row['Naam'].strip()}**] ({student_row['Opleiding']} / {student_row['Klas']}): Kans op uitval: {result['probability']:.1%}]"
            )

            with st.spinner(
                "Bezig met het ophalen van uitleg en feature importance...",
                show_time=True,
            ):
                # Fetch explanation and feature importance

                exp_response = requests.post(
                    "http://edflaptop.dyndns-home.com:8000/explain_risk",
                    json={
                        "student": student_row[features].to_dict(),
                        "prediction": result["prediction"],
                        "probability": result["probability"],
                    },
                    timeout=90,
                )

                st.info(exp_response.json()["explanation"])

                try:
                    fi_resp = requests.post(
                        "http://edflaptop.dyndns-home.com:8000/feature_importance",
                        json={"student": student_row[features].to_dict()},
                        timeout=30,
                    )
                    print(fi_resp)
                    fi_resp.raise_for_status()  # Check for HTTP errors
                    if fi_resp.status_code != 200:
                        st.error(
                            f"Failed to fetch feature importance: {fi_resp.status_code} - {fi_resp.text}"
                        )
                        return

                    fi = fi_resp.json()["feature_importance"]
                    fi_str = ", ".join([f"{k}: {v:.2f}" for k, v in fi.items()])
                    st.caption(
                        f"💡:red[**Belangrijkste risicofactoren (SHAP):**] **{fi_str.upper()}**"
                    )
                    st.write(tijd_verschil_microsec(start_tijd))

                except requests.exceptions.RequestException as e:
                    st.error(
                        f"An error occurred while fetching feature importance: {e}"
                    )
                    pass  # Replace 'continue' with 'pass' or handle the condition appropriately

    # risicostudenten = []
    teller = 0

    st.subheader(
        f"📊 :blue[**Voorspellen uitvalrisico studenten van de geselecteerde groep**]"
    )

    with st.expander(
        f""" :red[**Voorspel het uitvalrisico van studenten in de geselecteerde groep -**] :blue[**{opleiding.capitalize()}, {'alle klassen' if klas == 'Alle' else klas}, {'alle mentoren' if mentor == 'Alle' else mentor}**]""",
        expanded=True,
        icon="💡",
    ):

        st.markdown(
            """
                    *De gegevens van 10.000 **fictieve** studenten zijn gebruikt bij 
                    het trainen van het voorspelmodel (Random Forrest Classifier). 
                    Kenmerken: **Cijfer, Aanwezigheid, Waarschuwingen, EC (studiepunten)**,
                    voorspellen in het model de kans op een **1 (studieuitval)** of een 
                    **0 (geen uitval)**.*
                    """
        )
        
        @st.fragment()
        def start_prediction():
            """
            Start the prediction process for dropout risk.
            """

            # Start the prediction process
            with st.spinner("📊 :green[**Bezig met voorspellen...**]", show_time=True):

                placeholder = st.empty()
                teller_local = 0
                risicostudenten = []  # Initialize the list to store at-risk students

                for _, row in dff.iterrows():
                    teller_local += 1
                    student_dict = row[features].to_dict()
                    # st.write(student_dict)

                    pred_response = requests.post(
                        "http://edflaptop.dyndns-home.com:8000/predict_dropout",
                        json={"student": student_dict},
                        timeout=30,
                    )

                    result = pred_response.json()
                    # st.write(result)

                    if result["prediction"] == 1:
                        # Ensure risicostudenten is initialized before appending
                        # risicostudenten = [] # Initialize the list at the beginning of the function    or script
                        risicostudenten.append((row, result))
                        placeholder.empty()

                        placeholder.write(
                            f"""##### Studenten met verhoogd risico (:red[**{len(risicostudenten)}**])\n🔄 :blue[**Bezig met student {teller_local} van {len(dff)}...**]\n\n🔴 {teller_local}. :red[**{row['Naam'].strip()}**] ({row['Opleiding']} /  {row['Klas']})\n:red[**Kans op uitval:**]**{result['probability']:.1%}**\n\n:red[**{risicostudenten}**] """
                        )
                        time.sleep(0.1)  # Add a small delay to simulate processing time
                    else:
                        placeholder.write(
                            f"""##### Studenten met verhoogd risico (:red[**{len(risicostudenten)}**])\n🔄 :blue[**Bezig met student {teller_local} van {len(dff)}...**]\n\n✅ {teller_local}. **{row['Naam'].strip()}** **({row['Opleiding']} / {row['Klas']})**    :green[**lage kans op uitval**]\n\n:red[**{risicostudenten}**] """
                        )

            placeholder.empty()
            teller_local = 0
            # Check if risicostudenten is defined and not empty
            if not risicostudenten:
                risicostudenten = []
            if risicostudenten:
                for student in risicostudenten:
                    teller_local += 1
                    row, result = student
                    st.markdown(
                        f"""🔴 **{teller_local}.**  :red[**{row['Student-ID']}, {row['Naam'].strip()}, ({row['Opleiding']} / {row['Klas']}):  Kans op uitval: {result['probability']:.1%}**] """
                    )
                    row = ""
                    result = ""

                return details_uitval(risicostudenten)

            else:
                st.success("Geen risicostudenten in deze selectie!")

        if st.button(f"💡:blue[**Voorspel** ]"):

            st.write("----------------------------------------")
            start_prediction()
            st.write("----------------------------------------")

    @st.fragment()
    def toon_individuele_advies(opleiding: str = opleiding):
        """Toon individueel advies voor een geselecteerde student."""
        global dff
        global df
        if dff is None or dff.empty:
            return st.warning(
                "Geen data beschikbaar voor individuele advies. Zorg ervoor dat de data is geladen."
            )

        if dff is not None and not dff.empty:
            st.markdown("### 📚 :blue[**Management Advies bij een individueel geval**]")

            if True:
                student_data = pd.DataFrame()
                student_naam = st.selectbox(
                    "💡:red[**Selecteer Student Naam**]",
                    options=(["Maak een keuze"] + dff["Naam"].unique().tolist()),
                    format_func=lambda x: (
                        str(x) if x is not None else "Geen naam geselecteerd"
                    ),
                    index=0,
                )

                if student_naam and student_naam != "Maak een keuze":
                    student_id = (
                        dff[dff["Naam"] == student_naam]["Student-ID"].values[0]
                        if dff is not None
                        else None
                    )

                    if student_id is not None:
                        st.session_state.student_id = student_id
                        student_data = dff[dff["Student-ID"] == student_id]

                    if not student_data.empty:
                        student_data = student_data[
                            ["Naam"]
                            + VAKKEN
                            + ["Cijfer", "Aanwezigheid", "Waarschuwingen", "EC"]
                        ]

                        student_id = student_id if "student_id" in locals() else None  # type: ignore
                    if student_id is None:
                        student_id = None  # Initialize student_id with a default value
                    st.write(f"### Resultaten student: {student_id}, {student_naam}")
                    student_data = student_data.reset_index(drop=True)
                    st.session_state.student_naam = student_naam
                    st.session_state.student_id = student_id
                    st.session_state.opleiding = opleiding
                    st.session_state.klas = klas
                    st.session_state.mentor = mentor
                    st.dataframe(student_data, use_container_width=True, height=100)

                    if st.button(
                        f"💡:blue[**Genereer adviesrapport over student {st.session_state.get('student_id', 'Onbekend')} {st.session_state.get('student_naam', 'Onbekend')}**]"
                    ):

                        try:
                            start_tijd = set_starttijd()

                            with st.spinner(
                                "Bezig met het genereren van adviesrapport... ",
                                show_time=True,
                            ):

                                # Convert the student data to a dictionary payload;
                                # if there are multiple records, select the first one for the advi  ce    generation
                                payload = {"student": student_data.to_string()}
                                prompt = f"""
                        Je bent een onderwijsadviseur. 
                        Je geeft aan het management, en de mentor van student:
                        <student>
                        
                        {st.session_state.get('student_id', 'Onbekend')} {st.session_state.get('student_naam', 'Onbekend')},
                        
                        </student>
                        een gedegen, op feiten gebaseerd advies over 
                        de prestaties van de student en of de student een positief of negatief
                        Bindend Studie Advies (BSA) krijgt.: 
                        Het studentnummer is:
                        {st.session_state.get('student_id', 'Onbekend')}
                        De studentnaam is:
                        {st.session_state.get('student_naam', 'Onbekend')}
                         
                        De student is ingeschreven in de opleiding: 
                        {opleiding}, klas {klas}, en heeft als mentor {mentor}.
                        Je gaat nu een advies geven aan de hand van deze student gegevens:
                        {payload}
                        
                        Heeft de student voldoende kans om de opleiding af te 
                        kunnen ronden? Normering is gemiddeld Cijfer boven de 6, EC boven de 40, Aanwezigheid boven de 85%, en maximaal 1 Waarschuwing. Haalt de student de norm voor de BSA 
                        (Minstens 40 EC's)? En kan de student dan doorstromen 
                        naar het tweede jaar? Of is een overstap naar 
                        een andere opleiding beter? 
                        Zomaar wat vragen die een mentor en een manager zouden 
                        willen weten. Je gaat nu een fraai advies geven aan de 
                        hand van deze student gegevens:
                        {payload}
                        
                        De student is ingeschreven in de opleiding: 
                        {opleiding}, {klas} en  heeft als mentor {mentor}.
                        Richt je op de vakken: 
                        {', '.join(VAKKEN)}, en op de Aanwezigheid, EC,  
                        Waarschuwingen,en Cijfer. 
                        Deze zijn van belang voor het Bindend Studie Advies (BSA). 
                        De norm is op dit moment 40 EC's en een 
                        gemiddeld cijfer van 6.0, en Aanwezigheid boven de 85%. 
                        EC's zijn behaalde studiepunten en bepalen mede of 
                        een student door mag naar het tweede jaar. 
                        Geef een samenvatting van de prestaties van de student, 
                        benoem de sterke en zwakke punten, 
                        en geef een advies voor verbetering.
                        Geef een gedegen studieadvies voor een MBO-student.
                        Wees concreet en motiverend. 
                        Maak gebruik van het Markdown formaat en
                        antwoord altijd in het Nederlands
                      
                    
                        """

                                resp = requests.post(
                                    "http://edflaptop.dyndns-home.com:8000/studieadvies",
                                    headers={"Content-Type": "application/json"},
                                    json={"data": prompt},
                                    timeout=220,
                                )
                                if resp.status_code == 200:
                                    st.session_state.advies = resp.json().get(
                                        "advies", ""
                                    )
                                    st.write(tijd_verschil_microsec(start_tijd))

                                    st.markdown("### Adviesrapport")
                                    st.info(st.session_state.advies)

                                    return st.session_state.advies

                                # Check
                                if resp.status_code != 200:
                                    st.error(
                                        f"Fout bij het ophalen van het advies: {resp.status_code} - {resp.text}"
                                    )

                                    return resp

                                # antwoord = "
                                response = resp.json()["advies"]
                                print("Dit is de tweede")
                                print(response)
                                # return response

                                advies = response
                                if advies:
                                    if "advies" not in st.session_state:
                                        st.session_state.advies = advies
                                    else:
                                        st.session_state.advies += advies

                                    print(advies)

                                    # Add the response to the conversation
                                    st.markdown("### Adviesrapport")
                                    st.info(advies)
                                    st.session_state.advies = advies
                                    return advies

                                else:
                                    st.error("Geen advies ontvangen van de AI.")
                                    return

                        # advies_resp = requests.post(
                        #     "http://edflaptop.dyndns-home.com:8000/analyseer",
                        #     data=prompt,
                        #     timeout=60,
                        # )
                        # print(advies_resp)
                        # advies_resp.raise_for_status()  # Check for HTTPerrors
                        # if advies_resp.status_code != 200:
                        #     st.error(
                        #         f"Failed to fetch feature importance: {advies_resp.   status_code} - {advies_resp.text}"
                        #     )

                        except:
                            pass

        else:
            st.warning("Student niet gevonden.")

        with st.expander(
            ":red[**Bekijk Adviesrapport**]",
            expanded=False,
            icon="📄",
        ):
            if "student_id" not in st.session_state:
                st.session_state.student_id = None
            if "student_naam" not in st.session_state:
                st.session_state.student_naam = None

            student_id = st.session_state.get("student_id", None)
            student_naam = st.session_state.get("student_naam", None)
            if "advies" in st.session_state:
                st.markdown(f"### Adviesrapport {student_naam},  ({student_id})")

                st.info(st.session_state.advies)
                # Download button for the advice report
                st.download_button(
                    label="Download Adviesrapport",
                    data=str(st.session_state.get("advies", "")),
                    file_name=f"adviesrapport_student_{student_id}_{student_naam}.md",
                    mime="text/markdown",
                )

    advies = toon_individuele_advies(opleiding)
    if advies:
        st.markdown("### Adviesrapport")
        st.info(advies)

        # Download button for the advice report
        st.download_button(
            label="Download Adviesrapport",
            data=str(advies),
            file_name=f"adviesrapport_student_{st.session_state.get('student_id', 'Onbekend')}_{st.session_state.get('student_naam', 'Onbekend')}.md",
            mime="text/markdown",
        )

    # --- Sectie 3: Roosteroptimalisatie (Demo) ---
    st.header("⏱ Roosteroptimalisatie")
    if st.button("Genereer Voorbeeldrooster (AI)"):
        # Simpele demo - zou gekoppeld kunnen worden aan echte planningtools
        st.write(
            """**AI-geoptimaliseerd rooster:**
            - Maandag: Nederlands (09:00-11:00)
            - Dinsdag: Praktijkles (13:00-15:00)
            - Woensdag: Zelfstudie + AI-mentorgesprek
            - Tip: Verdeel theorielessen beter over de week"""
        )

        # --- Functies voor Predictive Analytics ---
        # @st.fragment()

    # @st.fragment()
    # def train_predictive_model(dataframe):
    #     """
    #     Train een voorspellend model voor studentrisico.
    #     """
    #     # Controleer of de benodigde kolommen bestaan

    #     required_columns = [
    #         "Cijfer",
    #         "Burgerschap",
    #         "Nederlands",
    #         "Project_KT2",
    #         "Rekenen",
    #         "Communicatie",
    #         "Klas",
    #     ]
    #     for col in required_columns:
    #         if col not in dataframe.columns:
    #             raise KeyError(f"Kolom '{col}' ontbreekt in de dataset.")

    #     # Maak een kopie van de dataframe om SettingWithCopy warnings te voorkomen
    #     df_copy = dataframe.copy()

    #     # Voorbeeld: Gebruik een Random Forest Classifier
    #     # Voorspel of een student risico loopt (cijfer < 5.5)
    #     df_copy["Risico"] = [0 for _ in range(len(df_copy))]
    #     df_copy["RisicoW"] = (
    #         df_copy["Waarschuwingen"].apply(lambda x: 1 if x > 1 else 0).copy()
    #     )
    #     df_copy["RisicoC"] = (
    #         df_copy["Cijfer"].apply(lambda x: 1 if x < 5.0 else 0).copy()
    #     )
    #     df_copy["RisicoE"] = df_copy["EC"].apply(lambda x: 1 if x < 35 else 0).copy()
    #     df_copy["Risico"] = [
    #         i + j + k
    #         for i, j, k in zip(
    #             df_copy["RisicoW"], df_copy["RisicoC"], df_copy["RisicoE"]
    #         )
    #     ]
    #     df_copy["Risico"] = df_copy["Risico"].apply(lambda x: 1 if x > 0 else 0).copy()

    #     X = pd.get_dummies(
    #         df_copy[
    #             [
    #                 "Burgerschap",
    #                 "Nederlands",
    #                 "Project_KT2",
    #                 "Rekenen",
    #                 "Communicatie",
    #                 "Klas",
    #             ]
    #         ]
    #     )  # Eenvoudige features
    #     y = df_copy["Risico"]

    #     X_train_inner, X_test_inner, y_train_inner, y_test_inner = train_test_split(
    #         X, y, test_size=0.2
    #     )
    #     clf = RandomForestClassifier()
    #     clf.fit(X_train_inner, y_train_inner)
    #     test_columns = X_test_inner.columns
    #     return clf, X_test_inner, y_test_inner, test_columns

    # @st.fragment()
    # def train():
    #     """
    #     Train een voorspellend model voor studentrisico.
    #     """
    #     # if st.button("Train het Predicatieve Model met Random Forrest"):
    #     student_id = st.selectbox(
    #         "Selecteer Student-ID",
    #         options=(
    #             sorted(
    #                 list(
    #                     df[df["Opleiding"] == opleiding]["Student-ID"].unique().tolist()
    #                 )
    #             )
    #             if df is not None
    #             else []
    #         ),
    #     )

    #     if student_id:
    #         st.session_state.student_id = str(student_id)
    #         st.session_state.student_id_predict = str(student_id)

    #     if (
    #         not st.session_state.student_id_predict
    #         or str(st.session_state.student_id_predict).strip() == "1500"
    #     ):
    #         st.warning("Voer eerst een geldige Student-ID in.")

    #     if str(st.session_state.student_id_predict) == "1500":
    #         st.warning(
    #             "Student-ID 1500 is een voorbeeldstudent. Gebruik een andere ID voor echte voorspellingen."
    #         )
    #         return

    #     with st.spinner("Bezig het model te trainen..."):
    #         trained_model, X_test, y_test, test_columns = train_predictive_model(df)
    #         y_pred = trained_model.predict(X_test)
    #         accuracy = accuracy_score(y_test, y_pred)
    #         # Toon resultaten
    #         st.success(f"Model getraind! Nauwkeurigheid: {accuracy:.0%}")
    #         # Confusion matrix
    #         st.subheader("Modelprestaties")
    #         cm = confusion_matrix(y_test, y_pred)
    #         conf_matrix_fig = px.imshow(
    #             cm,
    #             labels=dict(x="Voorspeld", y="Werkelijk"),
    #             text_auto=True,
    #             aspect="auto",
    #         )
    #         st.plotly_chart(conf_matrix_fig)
    #         # Sla het getrainde model en de testcolumns op in session_state
    #         st.session_state.model = trained_model
    #         st.session_state.X_testcolumns = test_columns

    # Voorspelling voor individuele student

    # st.subheader(":blue[**Risico-inschatting per Student**]")
    # with st.expander(
    #     f":blue[**Risicovoorspelling voor: {student_id}**]",
    #     expanded=False,
    #     icon="🔮",
    # ):

    #     # Input voor Student-ID
    #     if "student_id_predict" not in st.session_state:
    #         st.session_state.student_id_predict = "1500"
    #         student_id = "1500"
    #     else:
    #         student_id = st.session_state.student_id_predict

    #     student_id_predict = st.session_state.student_id_predict

    # st.selectbox(
    #     "Selecteer Student-ID",
    #     options=(
    #         df[df["Opleiding"] == opleiding]["Student-ID"].unique().tolist()
    #         if df is not None
    #         else []
    #     ),
    #     index=0,
    # )

    # #text_input("Voer Student-ID", value=student_id)

    # if student_id_predict:
    #     # st.session_state.student_id_predict = student_id_predict
    #     # Controleer of de Student-ID geldig is
    #     # if not student_id_predict.isdigit():
    #     #     st.error("Voer een geldige numerieke Student-ID in.")
    #     #     return

    #     if (
    #         str(student_id_predict).strip() == "1500"
    #         or str(student_id_predict).strip() == ""
    #     ):
    #         st.warning(
    #             "Student-ID 1500 is een voorbeeldstudent. Gebruik een andere ID voor echte voorspellingen."
    #         )
    #         return

    #     st.session_state.student_id_predict = student_id_predict
    # elif "student_id_predict" in st.session_state:
    #     student_id_predict = str(st.session_state.student_id_predict)

    # st.write(
    #     f":red[**Voorspelling voor Student-ID:**] :blue[**{student_id_predict}**]"
    # )
    # # Controleer of het model en de testcolumns zijn ingesteld
    # try:
    #     if df is not None:
    #         try:
    #             if not student_id_predict or not student_id_predict.strip():
    #                 raise ValueError("Student-ID is None or empty")
    #             sid = int(student_id_predict)
    #         except (TypeError, ValueError):
    #             st.error("Ongeldige of ontbrekende Student-ID")
    #             st.stop()

    #         student_record = df[df["Student-ID"] == sid]
    #         (
    #             st.write(
    #                 f"""Naam van student: :blue[**{student_record["Naam"].values[0]}**]"""
    #             )
    #             if not student_record.empty
    #             else "Student niet gevonden"
    #         )
    #     else:
    #         st.warning("Dataframe is leeg of niet geladen.")
    #         student_record = pd.DataFrame()  # Lege DataFrame als fallback
    #     if not student_record.empty:
    #         # Preprocess student data
    #         processed_data = pd.get_dummies(
    #             student_record[
    #                 [
    #                     "Burgerschap",
    #                     "Nederlands",
    #                     "Project_KT2",
    #                     "Rekenen",
    #                     "Communicatie",
    #                     "Klas",
    #                 ]
    #             ]
    #         )

    #         processed_data = processed_data.reindex(
    #             columns=st.session_state.X_testcolumns, fill_value=0
    #         )
    #         # Voorspelling
    #         risk = predict_student_risk(st.session_state.model, processed_data)
    #         st.write(
    #             f"Voorspeld risico voor student {st.session_state.student_id_predict}: {risk} "
    #         )
    #         # AI-uitleg
    #         explanation_prompt = f"""
    #         Leg uit waarom een student risico loopt in het Nederlands,
    #         gebruikmakend van deze data:
    #         {student_record.to_string()}

    #         """
    #         explanation = (
    #             client.chat.completions.create(
    #                 model=MODEL,
    #                 messages=[{"role": "user", "content": explanation_prompt}],
    #                 temperature=0.3,
    #             )
    #             .choices[0]
    #             .message.content
    #         )
    #         st.write("**AI-uitleg:**", explanation)
    #     else:
    #         st.warning("Student niet gevonden.")
    # except ValueError:
    #     st.error("Ongeldige Student-ID")

#     @st.fragment()
#     def predict_student_risk(model_used, student_features):
#         """Voorspel het risico voor een student op basis van zijn/haar kenmerken."""
#         if model_used is None:
#             return "Model is niet getraind. Train het model eerst."
#         if student_features is None or student_features.empty:
#             return "Geen geldige studentgegevens beschikbaar voor voorspelling."

#         # Zorg ervoor dat de features overeenkomen met het model
#         prediction = model_used.predict(student_features)

#         return "**Hoog risico** ❗" if prediction[0] == 1 else "**Laag risico** ✅"

#     # --- Aangepaste Dashboard Secties ---
#     def individu_voorspellen():
#         """
#         Voorspel het risico van studie-uitval voor een individuele student.
#         """
#         st.markdown(
#             "### Voorspel het risico van studie-uitval voor een individuele student"
#         )
#         st.write(
#             "Gebruik de AI om het risico van studie-uitval te voorspellen voor een individuele student op basis van zijn/haar gegevens."
#         )

#         st.subheader(
#             "🕵🏽 :blue[**Voorspellen van risico studieuitval bij 1 individuele student**]"
#         )
#         model = None
#         X_testcolumns = None
#         if df is not None and not df.empty:
#             train()


# if st.button(
#     "Train het model en voorspel het risico van studie-uitval voor een individuele student"
# ):
#     try:

#         if "individu_voorspellen" in globals():
#             if "individu_voorspellen" in globals():
#                 if "individu_voorspellen" in globals():
#                     individu_voorspellen()  # type: ignore
#                 else:
#                     st.error("Function 'individu_voorspellen' is not defined.")
#             else:
#                 st.error("Function 'individu_voorspellen' is not defined.")
#         else:
#             st.error("Function 'individu_voorspellen' is not defined.")
#     except:
#         pass

instructie_AI_coach = f"""
        Je bent een behulpzame en vriendelijke, Nederlands 
        sprekende managementadviseur. Je heet Karel van Oort, 
        en je ondersteunt de werkzaamheden van een manager in het onderwijs.
        Je hebt veel kennis van het onderwijs, 
        en je bent goed in het geven van advies en het beantwoorden van vragen.
        Je bent in staat om complexe vragen te beantwoorden 
        en je kunt goed omgaan met verschillende situaties.
        Je bent in staat om de manager te ondersteunen 
        bij het nemen van beslissingen en het oplossen van problemen.
        Je bent in staat om de manager te helpen 
        bij het maken van plannen en het uitvoeren van taken.
        Je bent in staat om de manager 
        te helpen bij het maken van rapportages en het analyseren van gegevens.
        Je hebt altijd deze data tot je beschikking:
        - Opleidingen: {', '.join(OPLEIDINGEN)}
        - Klassen: {', '.join(KLASSEN)}
        - Vakken: {', '.join(VAKKEN)}
        - Mentoren: {', '.join(MENTOREN)}
        - Dataframe met studentgegevens:
        <data>
            {(df.to_string(index=False) if df is not None else "Geen data beschikbaar")}
        </data>
        
        Beantwoord de vraag van de manager zo goed mogelijk, en geef een gedegen advies.
        Gebruik het Markdown formaat voor je antwoord en antwoord altijd in het Nederlands.
        Zorg ervoor dat je antwoord duidelijk en gestructureerd is, 
        en gebruik eventueel kopjes om de verschillende onderdelen van je antwoord te scheiden.
        Als je een vraag niet kunt beantwoorden, geef dan aan dat 
        je het antwoord niet weet, en bied aan om het antwoord later te zoeken.
        Zorg ervoor dat je antwoord altijd relevant is voor de vraag van de manager, 
        en dat je antwoord altijd gericht is op het helpen van de manager.
        
        """


@logger
def ai_chatbot(user_question):
    """
    AI Chatbot functie om vragen van een onderwijsmanager te beantwoorden.
    """
    global instructie_AI_coach
    global client

    if "managermessages" not in st.session_state:
        st.session_state.managermessages = []
    if not st.session_state.managermessages:
        st.session_state.managermessages = [
            {"role": "system", "content": instructie_AI_coach},
        ]

    managermessages = st.session_state.managermessages

    managermessages.append(
        {"role": "user", "content": f"Beantwoord de vraag: {user_question}"}
    )

    # container = client.containers.create(name="test-container")

    resp = requests.post(
        "http://edflaptop.dyndns-home.com:8000/chat",
        headers={"Content-Type": "application/json"},
        json={"data": managermessages},
        timeout=220,
    )
    if resp.status_code == 200:
        if "antwoord" not in st.session_state:
            st.session_state.antwoord = resp.json().get("antwoord", "")
        else:
            st.session_state.antwoord = resp.json().get("antwoord", "")

        with st.chat_message("assistant"):
            st.write(str(st.session_state.antwoord))

            return str(st.session_state.antwoord)

    # Check
    if resp.status_code != 200:
        st.error(
            f"Fout bij het ophalen van het antwoord: {resp.status_code} - {resp.text}"
        )

        return resp

    # Convert messages to the correct type for OpenAI API


chatbot_response = None
user_question = st.chat_input("Stel een vraag aan de AI manageradviseur:")
if user_question:
    with st.chat_message("user"):
        st.write(user_question)

    chatbot_response = ai_chatbot(user_question)

if chatbot_response and chatbot_response != "":

    with st.chat_message("assistant"):
        # st.write("🤖 AI Adviseur antwoord:")
        st.write(f" 🤖 {chatbot_response}")


with bottom():
    st.markdown(
        f"""
| :copyright: :grey[EdF 2025] | :grey[ed.de.feber@gmail.com] | :grey[{vandaag_str}] |
| :------ | :------ | :------ |
"""
    )
